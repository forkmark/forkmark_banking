"""Unit tests for core.compliance_reporter."""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import pytest

from core.compliance_reporter import ComplianceReporter, ValidationEvidence
from core.finance_evaluators import BiasDisparityEvaluator, NumericalFidelityEvaluator
from core.model_inventory import ModelInventory, ModelRecord, RiskTier
from core.regulatory_frameworks import ArtifactType, RegulatoryFramework
from core.statistical_analyzer import analyze
from core.store import Database


def _seed(tmp_path: Path) -> tuple[ComplianceReporter, ModelInventory]:
    db = Database(str(tmp_path / "memo.db"))
    inv = ModelInventory(db)
    inv.add_model(
        ModelRecord(
            model_id="credit-llm",
            display_name="Credit Decision Assistant",
            provider="anthropic",
            version="2.1",
            use_case="consumer credit adjudication",
            risk_tier=RiskTier.CRITICAL,
            regulatory_frameworks=[RegulatoryFramework.EU_AI_ACT],
            deployed_at=datetime(2025, 6, 1, tzinfo=timezone.utc),
            owner_team="Model Risk",
            present_artifacts=[
                ArtifactType.VALIDATION_MEMO.value,
                ArtifactType.TECHNICAL_DOCUMENTATION.value,
            ],
        )
    )
    return ComplianceReporter(db), inv


def _evidence() -> ValidationEvidence:
    stat = analyze(
        [0.9, 0.88, 0.91, 0.87, 0.93, 0.9, 0.92, 0.89],
        [0.2, 0.25, 0.22, 0.24, 0.21, 0.23, 0.19, 0.26],
    )
    bias = BiasDisparityEvaluator(threshold=1.2).evaluate(
        {"group_a": 0.9, "group_b": 0.5}  # fails 1.2x threshold
    )
    fidelity = NumericalFidelityEvaluator().evaluate(
        "Net income was $4.2 million.", "Net income was $9.9 million."
    )
    decisions = [
        {"choice": "A", "confidence": "high",
         "rationale_for_choice": "clearer numerical grounding",
         "rationale_for_rejection": "vague figures"},
        {"choice": "A", "confidence": "medium",
         "rationale_for_choice": "better numerical accuracy", "rationale_for_rejection": ""},
    ]
    return ValidationEvidence(
        statistical_results=[stat],
        bias_results=[bias],
        numerical_fidelity_results=[fidelity],
        decisions=decisions,
        evaluator_suite=["numerical_fidelity", "bias_disparity"],
    )


def test_memo_has_all_nine_sections(tmp_path: Path) -> None:
    reporter, _ = _seed(tmp_path)
    memo = reporter.generate_validation_memo(
        "credit-llm", RegulatoryFramework.EU_AI_ACT, _evidence()
    )
    for key in (
        "executive_summary",
        "scope_and_methodology",
        "statistical_results",
        "bias_and_fairness",
        "numerical_fidelity",
        "human_review_summary",
        "regulatory_mapping",
        "findings_and_recommendations",
        "sign_off",
    ):
        assert key in memo
    assert memo["executive_summary"]["model_name"] == "Credit Decision Assistant"
    assert memo["executive_summary"]["risk_tier"] == "CRITICAL"


def test_statistical_and_human_sections_populated(tmp_path: Path) -> None:
    reporter, _ = _seed(tmp_path)
    memo = reporter.generate_validation_memo(
        "credit-llm", RegulatoryFramework.EU_AI_ACT, _evidence()
    )
    assert "Win rate" in memo["statistical_results"][0]["plain_english"]
    hr = memo["human_review_summary"]
    assert hr["total_decisions"] == 2
    assert hr["choice_distribution"].get("A") == 2
    # Theme extraction picks up the recurring rationale keyword.
    assert any(t["theme"] == "numerical" for t in hr["rationale_themes"])


def test_findings_flag_bias_and_missing_artifacts(tmp_path: Path) -> None:
    reporter, _ = _seed(tmp_path)
    memo = reporter.generate_validation_memo(
        "credit-llm", RegulatoryFramework.EU_AI_ACT, _evidence()
    )
    findings = memo["findings_and_recommendations"]
    categories = {f["category"] for f in findings}
    assert "Bias & fairness" in categories
    assert "Numerical fidelity" in categories
    assert "Regulatory coverage" in categories
    # The failing bias check is HIGH severity.
    assert any(
        f["severity"] == "HIGH" and f["category"] == "Bias & fairness"
        for f in findings
    )
    # EU AI Act requires several artifacts we did not attach.
    assert memo["regulatory_mapping"]["missing_count"] > 0
    assert memo["regulatory_mapping"]["coverage_complete"] is False


def test_generate_validation_memo_docx_is_valid(tmp_path: Path) -> None:
    reporter, _ = _seed(tmp_path)
    out = str(tmp_path / "memo.docx")
    path = reporter.generate_validation_memo_docx(
        "credit-llm", RegulatoryFramework.EU_AI_ACT, _evidence(), output_path=out
    )
    assert Path(path).exists() and Path(path).stat().st_size > 0

    from docx import Document  # type: ignore[import-untyped]

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Model Validation Memorandum" in text
    assert "Credit Decision Assistant" in text
    assert "8. Findings & Recommendations" in text
    # Regulatory mapping + findings + sign-off tables are present.
    assert len(doc.tables) >= 3


def _is_arabic(text: str) -> bool:
    return any("؀" <= ch <= "ۿ" for ch in text)


def test_memo_json_carries_arabic_for_every_section(tmp_path: Path) -> None:
    """§4(b) disclosure applies to the whole memo, not just the executive
    summary — every section that has substantive content should carry a
    genuinely Arabic (not empty, not placeholder) counterpart in the JSON
    memo, since the JSON is the single source of truth for the .docx too."""
    reporter, _ = _seed(tmp_path)
    memo = reporter.generate_validation_memo(
        "credit-llm", RegulatoryFramework.EU_AI_ACT, _evidence()
    )

    assert _is_arabic(memo["scope_and_methodology"]["summary_ar"])
    assert _is_arabic(memo["statistical_results"][0]["plain_english_arabic"])
    assert _is_arabic(memo["bias_and_fairness"]["assessments"][0]["verdict_ar"])
    assert _is_arabic(memo["numerical_fidelity"]["assessments"][0]["verdict_ar"])
    assert _is_arabic(memo["human_review_summary"]["summary_ar"])
    for artifact in memo["regulatory_mapping"]["artifacts"]:
        assert _is_arabic(artifact["status_ar"])
    for finding in memo["findings_and_recommendations"]:
        assert _is_arabic(finding["description_ar"])
        assert _is_arabic(finding["recommendation_ar"])
        assert _is_arabic(finding["severity_ar"])
        assert _is_arabic(finding["category_ar"])


def test_memo_docx_body_is_bilingual_not_just_headings(tmp_path: Path) -> None:
    """Regression test for the gap this closes: the .docx used to carry Arabic
    only in the title, executive summary, and sign-off boilerplate, leaving
    sections 2-8 English-only. Every section's Arabic paragraph should now be
    present alongside its English content."""
    reporter, _ = _seed(tmp_path)
    out = str(tmp_path / "memo_bilingual.docx")
    path = reporter.generate_validation_memo_docx(
        "credit-llm", RegulatoryFramework.EU_AI_ACT, _evidence(), output_path=out
    )

    from docx import Document  # type: ignore[import-untyped]

    doc = Document(path)
    # Collect Arabic text from top-level paragraphs AND table cells (findings
    # and regulatory-mapping Arabic lives in cell paragraphs, not doc paragraphs).
    all_text = list(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.extend(p.text for p in cell.paragraphs)
    arabic_paragraphs = [t for t in all_text if _is_arabic(t)]

    # One Arabic paragraph per section (2-8) at minimum, plus the pre-existing
    # title/exec-summary/sign-off ones — comfortably more than the old ~4.
    assert len(arabic_paragraphs) >= 12
    # Spot-check specific section content actually made it into the document,
    # not just the section headings.
    joined = "\n".join(arabic_paragraphs)
    assert "الأطر التنظيمية" in joined  # section 2 (scope & methodology)
    assert "معدّل تفوّق" in joined  # section 3 (statistical results)
    assert "نسبة التفاوت" in joined  # section 4 (bias & fairness)
    assert "درجة الدقة" in joined  # section 5 (numerical fidelity)
    assert "إجمالي القرارات" in joined  # section 6 (human review)
    assert "متوفر" in joined or "غير متوفر" in joined  # section 7 (regulatory mapping)


def test_unknown_model_raises(tmp_path: Path) -> None:
    reporter, _ = _seed(tmp_path)
    with pytest.raises(KeyError):
        reporter.generate_validation_memo("nope", RegulatoryFramework.CBUAE_MMS)
