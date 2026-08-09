"""Model validation memo generation for ForkMark.

Assembles the evidence ForkMark holds for a model — statistical comparison
results, fairness and numerical-fidelity checks, human review decisions, and the
model-inventory record — into a structured model validation memorandum aligned to
a chosen regulatory framework. The memo is available both as a JSON-friendly dict
(for the API and UI) and as a professionally formatted ``.docx`` (for the
validator's evidence pack).

The memo follows a nine-section structure familiar to model risk teams:
    1. Executive Summary
    2. Scope & Methodology
    3. Statistical Results
    4. Bias & Fairness Assessment
    5. Numerical Fidelity
    6. Human Review Summary
    7. Regulatory Mapping
    8. Findings & Recommendations
    9. Sign-off
"""
from __future__ import annotations

import logging
import re
import tempfile
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Optional

from core.finance_evaluators import DisparityResult, NumericalFidelityResult
from core.model_inventory import ModelInventory, ModelRecord, SupportsStore
from core.regulatory_frameworks import (
    FrameworkRequirements,
    RegulatoryFramework,
    get_framework_requirements,
)
from core.statistical_analyzer import StatisticalResult

__all__ = [
    "ValidationEvidence",
    "Finding",
    "ComplianceReporter",
    "build_validation_memo",
]

logger = logging.getLogger("forkmark.compliance_reporter")

# Words ignored when extracting rationale themes from human review text.
_STOPWORDS = frozenset(
    """the a an and or but of to in on for with is are was were be been it this that
    as at by from than then so if not no yes better worse more less much very output
    outputs model branch a b response answer both because which while about""".split()
)


@dataclass(frozen=True)
class ValidationEvidence:
    """Evidence bundle assembled for a model validation memo.

    All fields are optional so a memo can be produced from whatever evidence is
    available; missing evidence surfaces as findings in Section 8.
    """

    statistical_results: list[StatisticalResult] = field(default_factory=list)
    bias_results: list[DisparityResult] = field(default_factory=list)
    numerical_fidelity_results: list[NumericalFidelityResult] = field(
        default_factory=list
    )
    decisions: list[dict[str, Any]] = field(default_factory=list)
    evaluator_suite: list[str] = field(default_factory=list)
    validation_period_start: Optional[datetime] = None
    validation_period_end: Optional[datetime] = None


_SEVERITY_AR = {"HIGH": "مرتفع", "MEDIUM": "متوسط", "LOW": "منخفض", "INFO": "معلوماتي"}

_CATEGORY_AR = {
    "Statistical significance": "الدلالة الإحصائية",
    "Statistical power": "القوة الإحصائية",
    "Bias & fairness": "التحيّز والإنصاف",
    "Numerical fidelity": "الدقة العددية",
    "Human oversight": "الإشراف البشري",
    "Regulatory coverage": "التغطية التنظيمية",
    "Revalidation": "إعادة التحقق",
    "Overall": "عام",
}


@dataclass(frozen=True)
class Finding:
    """A single validation finding with a recommended action.

    ``description``/``recommendation`` are English; ``description_ar``/
    ``recommendation_ar`` are the Arabic counterparts required for CBUAE's
    Arabic-and-English disclosure obligation (§4(b)). Both are populated from the
    same interpolated values in :func:`_generate_findings`, not machine-translated
    at render time, so the two languages never disagree on a number.
    """

    severity: str  # "HIGH" | "MEDIUM" | "LOW" | "INFO"
    category: str
    description: str
    recommendation: str
    description_ar: str = ""
    recommendation_ar: str = ""

    def to_dict(self) -> dict[str, str]:
        return {
            "severity": self.severity,
            "severity_ar": _SEVERITY_AR.get(self.severity, self.severity),
            "category": self.category,
            "category_ar": _CATEGORY_AR.get(self.category, self.category),
            "description": self.description,
            "description_ar": self.description_ar,
            "recommendation": self.recommendation,
            "recommendation_ar": self.recommendation_ar,
        }


def _summarize_decisions(decisions: list[dict[str, Any]]) -> dict[str, Any]:
    """Aggregate human review decisions: choice / confidence mix + rationale themes."""
    if not decisions:
        return {
            "total_decisions": 0,
            "choice_distribution": {},
            "confidence_distribution": {},
            "rationale_themes": [],
            "summary_ar": "لم يتم تسجيل أي قرارات مراجعة بشرية.",
        }
    choices = Counter(str(d.get("choice", "unknown")) for d in decisions)
    confidences = Counter(str(d.get("confidence", "unknown")) for d in decisions)

    words: Counter[str] = Counter()
    for d in decisions:
        text = f"{d.get('rationale_for_choice', '')} {d.get('rationale_for_rejection', '')}"
        for token in re.findall(r"[a-zA-Z]{4,}", text.lower()):
            if token not in _STOPWORDS:
                words[token] += 1
    themes = [{"theme": w, "count": n} for w, n in words.most_common(8)]
    themes_str = ", ".join(f"{t['theme']} ({t['count']})" for t in themes) or "لا يوجد"

    return {
        "total_decisions": len(decisions),
        "choice_distribution": dict(choices),
        "confidence_distribution": dict(confidences),
        "rationale_themes": themes,
        "summary_ar": (
            f"إجمالي القرارات: {len(decisions)}. توزيع الخيارات: {dict(choices)}. "
            f"توزيع مستوى الثقة: {dict(confidences)}. المحاور المتكررة في "
            f"المبررات: {themes_str}."
        ),
    }


def _generate_findings(
    model: ModelRecord,
    requirements: FrameworkRequirements,
    evidence: ValidationEvidence,
    missing_artifacts: list[str],
) -> list[Finding]:
    """Derive findings from failed thresholds, missing evidence, and due dates."""
    findings: list[Finding] = []

    # Statistical adequacy.
    for i, sr in enumerate(evidence.statistical_results, start=1):
        if not sr.is_significant:
            findings.append(
                Finding(
                    "MEDIUM",
                    "Statistical significance",
                    f"Comparison {i} is not statistically significant "
                    f"(adjusted p = {sr.adjusted_p_value:.3f}).",
                    "Increase the evaluation sample size or treat the branches as "
                    "equivalent for this metric.",
                    f"المقارنة رقم {i} غير ذات دلالة إحصائية "
                    f"(القيمة الاحتمالية المعدَّلة = {sr.adjusted_p_value:.3f}).",
                    "زيادة حجم عينة التقييم أو اعتبار الفرعين متكافئين بالنسبة لهذا "
                    "المقياس.",
                )
            )
        if abs(sr.effect_size) < sr.minimum_detectable_effect:
            findings.append(
                Finding(
                    "LOW",
                    "Statistical power",
                    f"Comparison {i} may be underpowered: |d|={abs(sr.effect_size):.2f} "
                    f"< MDE={sr.minimum_detectable_effect:.2f} at n={sr.sample_size}.",
                    "Expand the evaluation set to reach adequate power (>= 0.80).",
                    f"قد تكون المقارنة رقم {i} ذات قوة إحصائية غير كافية: "
                    f"|d|={abs(sr.effect_size):.2f} أقل من الحد الأدنى للتأثير "
                    f"القابل للاكتشاف={sr.minimum_detectable_effect:.2f} عند حجم "
                    f"عينة n={sr.sample_size}.",
                    "توسيع مجموعة التقييم للوصول إلى قوة إحصائية كافية (0.80 أو "
                    "أعلى).",
                )
            )

    # Bias / fairness.
    if requirements.bias_test_required and not evidence.bias_results:
        findings.append(
            Finding(
                "HIGH",
                "Bias & fairness",
                f"{requirements.name} requires bias testing, but no bias/fairness "
                "assessment was supplied.",
                "Run BiasDisparityEvaluator across the relevant demographic groups "
                "and attach the results.",
                f"يتطلب إطار {requirements.name} إجراء اختبار للتحيّز، إلا أنه لم "
                "يُقدَّم أي تقييم للتحيّز أو الإنصاف.",
                "تشغيل أداة تقييم التحيّز (BiasDisparityEvaluator) عبر الفئات "
                "الديموغرافية ذات الصلة وإرفاق النتائج.",
            )
        )
    for br in evidence.bias_results:
        if not br.passes_threshold:
            findings.append(
                Finding(
                    "HIGH",
                    "Bias & fairness",
                    f"Disparity ratio {br.disparity_ratio:.2f} exceeds the "
                    f"{br.threshold:.2f} threshold (max '{br.max_group}' vs "
                    f"min '{br.min_group}').",
                    "Investigate the driver of the disparity and remediate before "
                    "production use.",
                    f"نسبة التفاوت {br.disparity_ratio:.2f} تتجاوز الحد المسموح به "
                    f"{br.threshold:.2f} (الفئة الأعلى '{br.max_group}' مقابل الفئة "
                    f"الأدنى '{br.min_group}').",
                    "التحقيق في سبب التفاوت ومعالجته قبل الاستخدام في بيئة "
                    "الإنتاج.",
                )
            )

    # Numerical fidelity.
    for nf in evidence.numerical_fidelity_results:
        if not nf.is_faithful:
            findings.append(
                Finding(
                    "HIGH",
                    "Numerical fidelity",
                    f"{len(nf.flagged_numbers)} unsupported figure(s) detected in "
                    "model output relative to the source document.",
                    "Review flagged figures for material misstatement before relying "
                    "on the model's numeric output.",
                    f"تم رصد {len(nf.flagged_numbers)} رقم/أرقام غير مدعومة في "
                    "مخرجات النموذج مقارنةً بالمستند المصدر.",
                    "مراجعة الأرقام المرصودة للتحقق من عدم وجود أخطاء جوهرية قبل "
                    "الاعتماد على المخرجات العددية للنموذج.",
                )
            )

    # Human oversight.
    if requirements.human_oversight_required and not evidence.decisions:
        findings.append(
            Finding(
                "MEDIUM",
                "Human oversight",
                f"{requirements.name} requires documented human oversight, but no "
                "human review decisions were recorded.",
                "Capture human review decisions with rationale and confidence.",
                f"يتطلب إطار {requirements.name} إشرافاً بشرياً موثَّقاً، إلا أنه لم "
                "يتم تسجيل أي قرارات مراجعة بشرية.",
                "تسجيل قرارات المراجعة البشرية مع توضيح المبررات ومستوى الثقة.",
            )
        )

    # Regulatory artifact coverage.
    for artifact in missing_artifacts:
        findings.append(
            Finding(
                "MEDIUM",
                "Regulatory coverage",
                f"Required artifact '{artifact}' is not on file for this model.",
                f"Produce and attach the '{artifact}' artifact.",
                f"المستند المطلوب '{artifact}' غير موجود في ملف هذا النموذج.",
                f"إعداد وإرفاق المستند '{artifact}'.",
            )
        )

    # Revalidation currency.
    if model.next_validation_due is not None and model.next_validation_due < _now():
        due = model.next_validation_due.date().isoformat()
        findings.append(
            Finding(
                "MEDIUM",
                "Revalidation",
                f"Model is overdue for revalidation (due {due}).",
                "Schedule and complete a full revalidation.",
                f"تجاوز النموذج الموعد المحدد لإعادة التحقق (المستحق بتاريخ {due}).",
                "جدولة وإتمام عملية إعادة تحقق كاملة.",
            )
        )

    if not findings:
        findings.append(
            Finding(
                "INFO",
                "Overall",
                "No threshold breaches or missing artifacts were identified from the "
                "supplied evidence.",
                "Proceed to independent validator sign-off.",
                "لم يتم رصد أي تجاوز للحدود المسموح بها أو نقص في المستندات ضمن "
                "الأدلة المقدَّمة.",
                "المتابعة نحو اعتماد وتوقيع جهة تحقق مستقلة.",
            )
        )
    return findings


def _now() -> datetime:
    return datetime.now(timezone.utc)


def build_validation_memo(
    model: ModelRecord,
    framework: RegulatoryFramework,
    evidence: ValidationEvidence,
) -> dict[str, Any]:
    """Build the structured (JSON-friendly) validation memo dict.

    Pure function over a model record, a framework, and an evidence bundle — the
    single source of truth for both the JSON and ``.docx`` outputs.
    """
    requirements = get_framework_requirements(framework)
    present = set(model.present_artifacts)
    present_artifacts = [a for a in requirements.required_artifacts if a in present]
    missing_artifacts = [
        a for a in requirements.required_artifacts if a not in present
    ]

    period_start = (
        evidence.validation_period_start.date().isoformat()
        if evidence.validation_period_start
        else model.last_validated_at.date().isoformat()
        if model.last_validated_at
        else None
    )
    period_end = (
        evidence.validation_period_end.date().isoformat()
        if evidence.validation_period_end
        else _now().date().isoformat()
    )

    sample_sizes = [sr.sample_size for sr in evidence.statistical_results]
    findings = _generate_findings(model, requirements, evidence, missing_artifacts)

    return {
        "title": "ForkMark Model Validation Memorandum",
        "generated_at": _now().isoformat(),
        "executive_summary": {
            "model_id": model.model_id,
            "model_name": model.display_name,
            "provider": model.provider,
            "version": model.version,
            "use_case": model.use_case,
            "risk_tier": model.risk_tier.value,
            "framework": requirements.name,
            "jurisdiction": requirements.jurisdiction,
            "validation_period": {"start": period_start, "end": period_end},
            "status": model.status.value,
        },
        "scope_and_methodology": {
            "frameworks_assessed": [requirements.name],
            "reference": requirements.reference,
            "sample_sizes": sample_sizes,
            "total_comparisons": len(evidence.statistical_results),
            "evaluator_suite": list(evidence.evaluator_suite),
            "summary_ar": (
                f"الأطر التنظيمية المشمولة بالتقييم: {requirements.name} "
                f"({requirements.reference}). عدد المقارنات التي جرى تحليلها: "
                f"{len(evidence.statistical_results)}؛ أحجام العينات: "
                f"{sample_sizes or 'لا يوجد'}. مجموعة أدوات التقييم المستخدمة: "
                f"{', '.join(evidence.evaluator_suite) or 'لا يوجد'}."
            ),
        },
        "statistical_results": [
            {
                "win_rate": sr.win_rate,
                "confidence_interval_95": [sr.ci_lower, sr.ci_upper],
                "p_value": sr.p_value,
                "adjusted_p_value": sr.adjusted_p_value,
                "effect_size_cohens_d": sr.effect_size,
                "is_significant": sr.is_significant,
                "sample_size": sr.sample_size,
                "minimum_detectable_effect": sr.minimum_detectable_effect,
                "test_method": sr.method,
                "plain_english": sr.as_plain_english(),
                "plain_english_arabic": sr.as_plain_english_arabic(),
            }
            for sr in evidence.statistical_results
        ],
        "bias_and_fairness": {
            "required": requirements.bias_test_required,
            "assessments": [
                {
                    "per_group_scores": br.per_group_scores,
                    "disparity_ratio": br.disparity_ratio,
                    "threshold": br.threshold,
                    "passes_threshold": br.passes_threshold,
                    "verdict_ar": "مطابق" if br.passes_threshold else "غير مطابق",
                    "max_group": br.max_group,
                    "min_group": br.min_group,
                }
                for br in evidence.bias_results
            ],
            "overall_pass": all(br.passes_threshold for br in evidence.bias_results)
            if evidence.bias_results
            else None,
        },
        "numerical_fidelity": {
            "applicable": bool(evidence.numerical_fidelity_results),
            "assessments": [
                {
                    "score": nf.score,
                    "is_faithful": nf.is_faithful,
                    "verdict_ar": "مطابق" if nf.is_faithful else "غير مطابق",
                    "total_output_numbers": nf.total_output_numbers,
                    "flagged_numbers": [
                        {"value": f.value, "raw": f.raw, "kind": f.kind,
                         "reason": f.reason}
                        for f in nf.flagged_numbers
                    ],
                }
                for nf in evidence.numerical_fidelity_results
            ],
            "overall_pass": all(
                nf.is_faithful for nf in evidence.numerical_fidelity_results
            )
            if evidence.numerical_fidelity_results
            else None,
        },
        "human_review_summary": _summarize_decisions(evidence.decisions),
        "regulatory_mapping": {
            "framework": requirements.name,
            "artifacts": [
                {
                    "artifact": a,
                    "status": "PRESENT" if a in present else "MISSING",
                    "status_ar": "متوفر" if a in present else "غير متوفر",
                }
                for a in requirements.required_artifacts
            ],
            "present_count": len(present_artifacts),
            "missing_count": len(missing_artifacts),
            "coverage_complete": not missing_artifacts,
        },
        "findings_and_recommendations": [f.to_dict() for f in findings],
        "sign_off": {
            "validator_name": "",
            "validator_title": "",
            "validation_date": "",
            "approval_status": "",
            "comments": "",
        },
    }


class ComplianceReporter:
    """Generates model validation memos (JSON and .docx) from ForkMark evidence.

    Args:
        db: A ForkMark ``Database`` (or anything satisfying ``SupportsStore``) used
            to resolve the model-inventory record for a ``model_id``.
    """

    def __init__(self, db: SupportsStore) -> None:
        self._inventory = ModelInventory(db)

    def _resolve_model(self, model_id: str) -> ModelRecord:
        model = self._inventory.get_model(model_id)
        if model is None:
            raise KeyError(f"model not found in inventory: {model_id!r}")
        return model

    def generate_validation_memo(
        self,
        model_id: str,
        framework: RegulatoryFramework,
        evidence: Optional[ValidationEvidence] = None,
    ) -> dict[str, Any]:
        """Generate the structured validation memo dict for a model + framework."""
        model = self._resolve_model(model_id)
        memo = build_validation_memo(model, framework, evidence or ValidationEvidence())
        logger.info(
            "compliance: generated validation memo for %s under %s",
            model_id,
            framework.value,
        )
        return memo

    def generate_validation_memo_docx(
        self,
        model_id: str,
        framework: RegulatoryFramework,
        evidence: Optional[ValidationEvidence] = None,
        output_path: Optional[str] = None,
    ) -> str:
        """Generate the validation memo as a formatted ``.docx`` file.

        Returns:
            The path to the written ``.docx`` file.
        """
        memo = self.generate_validation_memo(model_id, framework, evidence)
        path = output_path or tempfile.mkstemp(
            suffix=".docx", prefix=f"validation_memo_{model_id}_"
        )[1]
        _render_docx(memo, path)
        logger.info("compliance: wrote validation memo docx to %s", path)
        return path


def _render_docx(memo: dict[str, Any], output_path: str) -> None:
    """Render a memo dict to a professionally structured .docx file."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    doc = Document()

    # ── Bilingual (Arabic RTL) helper ────────────────────────────────────────
    # CBUAE requires AI-decision disclosure/explainability in Arabic and English
    # (§4(b)), so every section of the memo — not just the title and executive
    # summary — carries its substantive content in both languages.
    def _style_rtl(p, text: str, *, bold: bool = False, size: int = 11):
        """Apply RTL paragraph direction + an Arabic-shaped run to a paragraph.

        Works on both top-level document paragraphs and table-cell paragraphs,
        since both are ``Paragraph`` objects with the same XML surface.
        """
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run._element.get_or_add_rPr().append(OxmlElement("w:rtl"))
        return p

    def _ar(text: str, *, bold: bool = False, size: int = 11):
        return _style_rtl(doc.add_paragraph(), text, bold=bold, size=size)

    def _ar_cell(cell, text: str, *, bold: bool = False, size: int = 9.5):
        """Add an Arabic RTL paragraph below a table cell's existing English text."""
        return _style_rtl(cell.add_paragraph(), text, bold=bold, size=size)

    brand = doc.add_paragraph()
    brand_run = brand.add_run("ForkMark")
    brand_run.bold = True
    brand_run.font.size = Pt(14)
    brand_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    title = doc.add_heading("Model Validation Memorandum", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    exec_summary = memo["executive_summary"]
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        f"{exec_summary['model_name']} ({exec_summary['version']}) — "
        f"{exec_summary['framework']}"
    ).italic = True
    _ar("مذكرة التحقق من صحة النموذج", bold=True, size=13)
    doc.add_paragraph(f"Generated: {memo['generated_at']}")

    # 1. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    _ar("١. الملخص التنفيذي", bold=True, size=11)
    period = exec_summary["validation_period"]
    for label, value in (
        ("Model", f"{exec_summary['model_name']} ({exec_summary['model_id']})"),
        ("Provider / Version", f"{exec_summary['provider']} / {exec_summary['version']}"),
        ("Use Case", exec_summary["use_case"]),
        ("Risk Tier", exec_summary["risk_tier"]),
        ("Framework", f"{exec_summary['framework']} — {exec_summary['jurisdiction']}"),
        ("Validation Period", f"{period['start']} to {period['end']}"),
        ("Inventory Status", exec_summary["status"]),
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    # Bilingual executive summary (Arabic) — CBUAE explainability expectation.
    _n_findings = len(memo["findings_and_recommendations"])
    _ar("الملخص التنفيذي", bold=True, size=12)
    _ar(
        f"يعرض هذا التقرير التحقق المستقل من النموذج «{exec_summary['model_name']}» "
        f"({exec_summary['model_id']})، المصنّف ضمن فئة المخاطر {exec_summary['risk_tier']}، "
        f"وفقاً لإطار {exec_summary['framework']}. خلص التحقق إلى "
        f"{_n_findings} ملاحظة/ملاحظات تتطلب المتابعة قبل الاعتماد النهائي من جهة تحقق مستقلة. "
        f"أُجري التحقق ذاتياً داخل بيئة البنك دون خروج أي بيانات للعملاء."
    )

    # 2. Scope & Methodology
    scope = memo["scope_and_methodology"]
    doc.add_heading("2. Scope & Methodology", level=1)
    _ar("٢. النطاق والمنهجية", bold=True, size=11)
    doc.add_paragraph(
        f"Frameworks assessed: {', '.join(scope['frameworks_assessed'])} "
        f"({scope['reference']})."
    )
    doc.add_paragraph(
        f"Comparisons analysed: {scope['total_comparisons']}; "
        f"sample sizes: {scope['sample_sizes'] or 'n/a'}."
    )
    doc.add_paragraph(
        f"Evaluator suite: {', '.join(scope['evaluator_suite']) or 'n/a'}."
    )
    _ar(scope["summary_ar"])

    # 3. Statistical Results
    doc.add_heading("3. Statistical Results", level=1)
    _ar("٣. النتائج الإحصائية", bold=True, size=11)
    if memo["statistical_results"]:
        for i, sr in enumerate(memo["statistical_results"], start=1):
            p = doc.add_paragraph(style="List Number")
            p.add_run(sr["plain_english"])
            _ar(sr["plain_english_arabic"], size=10)
    else:
        doc.add_paragraph("No statistical comparison results were supplied.")
        _ar("لم تُقدَّم أي نتائج مقارنة إحصائية.")

    # 4. Bias & Fairness Assessment
    doc.add_heading("4. Bias & Fairness Assessment", level=1)
    _ar("٤. تقييم التحيّز والإنصاف", bold=True, size=11)
    bias = memo["bias_and_fairness"]
    if bias["assessments"]:
        for b in bias["assessments"]:
            verdict = "PASS" if b["passes_threshold"] else "FAIL"
            doc.add_paragraph(
                f"[{verdict}] Disparity ratio {b['disparity_ratio']:.3f} "
                f"(threshold {b['threshold']:.2f}); max '{b['max_group']}' vs "
                f"min '{b['min_group']}'.",
                style="List Bullet",
            )
            _ar(
                f"[{b['verdict_ar']}] نسبة التفاوت {b['disparity_ratio']:.3f} "
                f"(الحد المسموح به {b['threshold']:.2f})؛ الفئة الأعلى "
                f"'{b['max_group']}' مقابل الفئة الأدنى '{b['min_group']}'.",
                size=10,
            )
    else:
        note = (
            "Bias testing is required for this framework but no assessment was supplied."
            if bias["required"]
            else "No bias/fairness assessment supplied."
        )
        doc.add_paragraph(note)
        note_ar = (
            "يتطلب هذا الإطار إجراء اختبار للتحيّز، ولم يتم تقديم أي تقييم."
            if bias["required"]
            else "لم يتم تقديم أي تقييم للتحيّز أو الإنصاف."
        )
        _ar(note_ar)

    # 5. Numerical Fidelity
    doc.add_heading("5. Numerical Fidelity", level=1)
    _ar("٥. الدقة العددية", bold=True, size=11)
    nf = memo["numerical_fidelity"]
    if nf["applicable"]:
        for a in nf["assessments"]:
            verdict = "PASS" if a["is_faithful"] else "FAIL"
            doc.add_paragraph(
                f"[{verdict}] Fidelity score {a['score']:.2f}; "
                f"{len(a['flagged_numbers'])} flagged of "
                f"{a['total_output_numbers']} figures.",
                style="List Bullet",
            )
            _ar(
                f"[{a['verdict_ar']}] درجة الدقة {a['score']:.2f}؛ تم رصد "
                f"{len(a['flagged_numbers'])} من أصل {a['total_output_numbers']} "
                f"رقماً.",
                size=10,
            )
    else:
        doc.add_paragraph("Not applicable for this model's use case.")
        _ar("لا ينطبق على حالة استخدام هذا النموذج.")

    # 6. Human Review Summary
    doc.add_heading("6. Human Review Summary", level=1)
    _ar("٦. ملخص المراجعة البشرية", bold=True, size=11)
    hr = memo["human_review_summary"]
    doc.add_paragraph(f"Total decisions: {hr['total_decisions']}.")
    if hr["total_decisions"]:
        doc.add_paragraph(f"Choice distribution: {hr['choice_distribution']}.")
        doc.add_paragraph(
            f"Confidence distribution: {hr['confidence_distribution']}."
        )
        themes = ", ".join(f"{t['theme']} ({t['count']})" for t in hr["rationale_themes"])
        doc.add_paragraph(f"Rationale themes: {themes or 'n/a'}.")
    _ar(hr["summary_ar"])

    # 7. Regulatory Mapping
    doc.add_heading("7. Regulatory Mapping", level=1)
    _ar("٧. الربط التنظيمي", bold=True, size=11)
    mapping = memo["regulatory_mapping"]
    map_table = doc.add_table(rows=1, cols=2)
    map_table.style = "Light Grid Accent 1"
    hdr = map_table.rows[0].cells
    hdr[0].text = "Required Artifact / المستند المطلوب"
    hdr[1].text = "Status / الحالة"
    for item in mapping["artifacts"]:
        row = map_table.add_row().cells
        row[0].text = item["artifact"]
        row[1].text = item["status"]
        _ar_cell(row[1], item["status_ar"])

    # 8. Findings & Recommendations
    doc.add_heading("8. Findings & Recommendations", level=1)
    _ar("٨. النتائج والتوصيات", bold=True, size=11)
    find_table = doc.add_table(rows=1, cols=4)
    find_table.style = "Light Grid Accent 1"
    fh = find_table.rows[0].cells
    fh[0].text, fh[1].text, fh[2].text, fh[3].text = (
        "Severity / الخطورة",
        "Category / الفئة",
        "Finding / الملاحظة",
        "Recommendation / التوصية",
    )
    for finding in memo["findings_and_recommendations"]:
        row = find_table.add_row().cells
        row[0].text = finding["severity"]
        _ar_cell(row[0], finding["severity_ar"])
        row[1].text = finding["category"]
        _ar_cell(row[1], finding["category_ar"])
        row[2].text = finding["description"]
        _ar_cell(row[2], finding["description_ar"])
        row[3].text = finding["recommendation"]
        _ar_cell(row[3], finding["recommendation_ar"])

    # 9. Sign-off
    doc.add_heading("9. Sign-off", level=1)
    _ar("٩. الاعتماد والتوقيع", bold=True, size=11)
    doc.add_paragraph(
        "This validation is not effective until signed by an independent validator."
    )
    _ar("الاعتماد والتوقيع", bold=True, size=12)
    _ar("لا يُعدّ هذا التحقق نافذاً إلا بعد توقيع جهة تحقق مستقلة.")
    sign_table = doc.add_table(rows=4, cols=2)
    sign_table.style = "Table Grid"
    labels = ["Independent Validator", "Title", "Approval Status", "Date"]
    for idx, label in enumerate(labels):
        sign_table.rows[idx].cells[0].text = label
        sign_table.rows[idx].cells[1].text = ""

    doc.save(output_path)
